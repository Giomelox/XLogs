from app.back_end.funções.funções_conectar_email import aux_path_XML_destino, download_dir, conectar_email_e_baixar_arquivos_HP, conectar_email_e_baixar_arquivos_Dell
from app.back_end.funções.funções_Gerais import PlanilhaManager, obter_configs, navigate
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium.webdriver.common.by import By
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from selenium import webdriver
from docx.oxml.ns import qn
from xml.dom import minidom
from docx import Document
import flet as ft
import shutil
import time
import os


def comparar_cprod(xml_file, cprod_excel, nNF_excel):
    try:
        tree = minidom.parse(xml_file)
        nf_list = tree.getElementsByTagName('nNF')  # Lista com os números de nNF
        det_list = tree.getElementsByTagName('det')  # Lista com os números de det

        # Verifica se há nNF correspondente
        for nf in nf_list:
            if nf.firstChild and nf.firstChild.data.strip() == nNF_excel.strip():  # Compara nNF
                # Para o nNF correspondente, procura os elementos <det>
                for det in det_list:
                    prod_element = det.getElementsByTagName('prod')[0] if det.getElementsByTagName('prod') else None  # Verifica se <prod> existe
                    if prod_element:
                        cProd = prod_element.getElementsByTagName('cProd')[0].firstChild.data.strip() if prod_element.getElementsByTagName('cProd') else None
                        if cProd and cProd == cprod_excel.strip():  # Compara cProd do XML com o valor do Excel
                            return det  # Retorna o <det> correspondente
        return None  # Retorna None se não houver correspondência
    except Exception as e:
        return None

# Botoes_entrada_Dell e HP
def biparxml(log_instance, page, views, current_view):

    planilha_manager = PlanilhaManager()

    navigate(current_view = current_view, page = page, views = views, view_name = 'logs e informações recebimento')

    botao_voltar = views['logs e informações recebimento'].content.controls[1].content.controls[0].controls[0]
    botao_voltar.disabled = True
    botao_voltar.content.color = ft.Colors.GREY_400
    page.update()

    log_instance.log_message('Iniciando inserção das chaves NFe...')

    usuario_elogistica = obter_configs().get('usuario_Elogistic')
    senha_elogistica = obter_configs().get('senha_Elogistic')
    
    substrings_bipar = []

    if planilha_manager.is_planilha_atualizada():
        planilha_manager.recarregar_planilha()
    
    planilha_df = planilha_manager.get_planilha()

    driver = webdriver.Chrome()
    
    try:
        for index, row in planilha_df.iterrows():

            cell_bipar = row[1]

            if isinstance(cell_bipar, str) and len(cell_bipar) >= 34:
                # Extrai a substring e adiciona à lista
                substring_bipar = cell_bipar
                substrings_bipar.append(substring_bipar)
            else:
                substrings_bipar.append(cell_bipar)
    except:
        driver.quit()

        log_instance.log_message('Erro: Verifique se alguma planilha foi selecionada')

        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE
        page.update()

    idx = 1

    # Acessa a página desejada
    driver.get('https://innovation.uninet.com.br/Ulog/Main.asp?All=reset')

    time.sleep(5)

    try:
        element_usuario = WebDriverWait(driver, 500).until(
            EC.presence_of_element_located((By.NAME, 'login'))
        )
        element_usuario.send_keys(usuario_elogistica)

        element_senha = WebDriverWait(driver, 2).until(
            EC.presence_of_element_located((By.NAME, 'password'))
        )
        element_senha.send_keys(senha_elogistica)

        element_submit = WebDriverWait(driver, 2).until(
            EC.presence_of_element_located((By.NAME, 'submit'))
        )
        element_submit.click()

        time.sleep(5)

        element_botao_notificar_recebimento = WebDriverWait(driver, 500).until(
            EC.presence_of_element_located((By.XPATH, "(//input[@id='submit'])[2]"))
        )
        element_botao_notificar_recebimento.click()
    
    except AssertionError:
        pass

    except:
        log_instance.log_message('Elemento não apareceu dentro do tempo esperado.')
        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE
        page.update()

    time.sleep(10)

    try:
        for x in substrings_bipar:

            element_xml = WebDriverWait(driver, 500).until(
                EC.presence_of_element_located((By.NAME, 'keynfe'))
            )
            element_xml.send_keys(x)

            time.sleep(2)

            element_submit_xml = driver.find_element(By.NAME, 'submit')
            element_submit_xml.click()
            
            log_instance.log_message(f'{idx} - Arquivo {x} processado')

            idx += 1

            time.sleep(5)

    except AssertionError:
        pass
            
    except Exception as e:
        log_instance.log_message(f'Não foi possível adicionar o arquivo {x}: {e}')
        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE

    time.sleep(5)

    substrings_bipar.clear()

    # Fechar o navegador
    driver.quit()

    log_instance.log_message('Inserção das chaves NFe concluída.')

    # Reabilitar o botão "Voltar" após a conclusão
    botao_voltar.disabled = False
    botao_voltar.content.color = ft.Colors.WHITE
    page.update()

# Botoes_devolução_Dell
def valores_devolução_DELL(page: ft.Page, log_instance, views, current_view):

    def mostrar_confirmacao(page):
        dialog.open = True
        page.update()

    # Função para o botão "Sim"
    def continuar(e):

        planilha_manager = PlanilhaManager()

        navigate(current_view = current_view, page = page, views = views, view_name = 'logs e informações devolução')

        botao_voltar = views['logs e informações devolução'].content.controls[1].content.controls[0].controls[0]
        botao_voltar.disabled = True
        botao_voltar.content.color = ft.Colors.GREY_400
        dialog.open = False
        page.update()

        usuario_elogistica = obter_configs().get('usuario_Elogistic')
        senha_elogistica = obter_configs().get('senha_Elogistic')

        # Lista todos os arquivos na pasta
        arquivos_excluir = os.listdir(aux_path_XML_destino)

        log_instance.log_message('Iniciando inserção de valores...')

        for arquivo_excluir in arquivos_excluir:
            caminho_arquivo_excluir = os.path.join(aux_path_XML_destino, arquivo_excluir)
            if os.path.isfile(caminho_arquivo_excluir):
                os.remove(caminho_arquivo_excluir)

        try:
            conectar_email_e_baixar_arquivos_Dell(log_instance)
        except Exception as e:
            log_instance.log_message(f'Erro: Verifique se alguma planilha foi selecionada. Feche o navegador\n{e}')
            botao_voltar.disabled = False
            botao_voltar.content.color = ft.Colors.WHITE
            page.update()
            return
        
        driver = webdriver.Chrome()

        # Acessa a página desejada
        driver.get('https://innovation.uninet.com.br/Ulog/Main.asp?All=reset')

        try:
            element_usuario = WebDriverWait(driver, 200).until(
                EC.presence_of_element_located((By.NAME, 'login'))
            )
            element_usuario.send_keys(usuario_elogistica)

            element_senha = driver.find_element(By.NAME, 'password')
            element_senha.send_keys(senha_elogistica)

            element_submit = driver.find_element(By.NAME, 'submit')
            element_submit.click()

            time.sleep(5)

            element_peças_pendentes_devolução = WebDriverWait(driver, 200).until(
                EC.presence_of_element_located((By.XPATH, '(//input[@id="submit"])[4]'))
            )
            element_peças_pendentes_devolução.click()

        except AssertionError:
            pass

        except:
            log_instance.log_message('O site demorou mais que o esperado para carregar.')
            botao_voltar.disabled = False
            botao_voltar.content.color = ft.Colors.WHITE
            page.update()

        idx_planilha = 1

        resultados = {}

        if planilha_manager.is_planilha_atualizada():
            planilha_manager.recarregar_planilha()
        
        planilha_df = planilha_manager.get_planilha()
   
        try:
            for index, row in planilha_df.iloc[:].iterrows():
                CHAMADO = str(row[0]).replace('.0', '').strip()
                NF = str(row[1]).replace('.0', '').strip()
                PN = str(row[2]).strip()
                PPID = str(row[4]).strip()
                STATUS_planilha = str(row[10]).strip()

                if STATUS_planilha.upper() == 'DEFECTIVE' or STATUS_planilha.upper() == 'DOA':

                    if PPID.upper() == 'X':
                        log_instance.log_message(f'{idx_planilha}: Peça sem PPID na planilha.')

                        try:
                            # Encontra o valor do produto
                            element_Vprod = elements.find_element(By.XPATH, './/td[9]')
                            element_Vprod_value = element_Vprod.text

                            if not os.path.exists(aux_path_XML_destino):
                                continue
                            else:  
                                # Lista todos os arquivos na pasta
                                arquivos = os.listdir(aux_path_XML_destino)

                                for arquivo in arquivos:
                                    caminho_arquivo = os.path.join(aux_path_XML_destino, arquivo)

                                    if os.path.isfile(caminho_arquivo):
                                        try:
                                            # Abrir o arquivo XML no modo leitura
                                            with open(caminho_arquivo, 'r'):

                                                # Comparar cProd do arquivo XML com o valor PN do Excel
                                                det_element = comparar_cprod(caminho_arquivo, PN, NF)

                                                if det_element:
                                                    # Buscar o valor de <orig> e <NCM> dentro do mesmo <det>
                                                    CST = det_element.getElementsByTagName('orig')
                                                    ICMS = det_element.getElementsByTagName('pICMS')
                                                    NCM_prod = det_element.getElementsByTagName('NCM')
                                                    Quantidade_prod = det_element.getElementsByTagName('qCom')
                                                    
                                                    chave_resultado = f"{PN}_{STATUS_planilha}_{NF}"

                                                    if CST and NCM_prod:
                                                        resultados[chave_resultado] = {
                                                            'VALOR': element_Vprod_value if element_Vprod else None,
                                                            'CST': str(CST[0].firstChild.data) if CST else None,
                                                            'ICMS': float(ICMS[0].firstChild.data) if ICMS else None,
                                                            'NCM': str(NCM_prod[0].firstChild.data) if NCM_prod else None,
                                                            'QTD.': float(Quantidade_prod[0].firstChild.data) if Quantidade_prod else None
                                                        }

                                                    else:
                                                        log_instance.log_message(f'Elementos CST, NCM e/ou Quantidade não encontrados no arquivo {arquivo}.')

                                        except AssertionError:
                                            pass

                                        except Exception as e:
                                            log_instance.log_message(f'Erro ao processar o arquivo {arquivo}: {e}')
                            idx_planilha += 1

                        except AssertionError:
                            pass           

                        except Exception as e:
                            log_instance.log_message(f'{idx_planilha}: Chamado {CHAMADO}/{PN} não consta no site.')
                            idx_planilha += 1
                            continue

                    else:
                        try:
                            
                            # Tenta encontrar o elemento correspondente
                            elements = WebDriverWait(driver, 20).until(
                                EC.presence_of_element_located((By.XPATH, f'//tr[contains(td[4], "{CHAMADO}") and contains(td[6], "{PN}")]'))
                            )

                            time.sleep(1)

                            if elements:
                                try:
                                    # Encontra o valor do produto
                                    element_Vprod = elements.find_element(By.XPATH, './/td[9]')
                                    element_Vprod_value = element_Vprod.text

                                    time.sleep(1)

                                    # Encontra o link <a> na mesma linha
                                    element_link = elements.find_element(By.XPATH, './/td[5]')
                                    element_link.click()

                                    time.sleep(1)
                                    
                                    # Insere o PPID da planilha no site
                                    inserir_ppid = WebDriverWait(driver, 60).until(
                                    EC.presence_of_element_located((By.NAME, 'PPID'))
                                    )
                                    inserir_ppid.send_keys(PPID)

                                    botao_enviar_ppid = WebDriverWait(driver, 20).until(
                                    EC.presence_of_element_located((By.NAME, 'submit'))
                                    )
                                    botao_enviar_ppid.click()

                                    time.sleep(1)

                                    alerta = driver.switch_to.alert
                                    alerta.accept()

                                    if not os.path.exists(aux_path_XML_destino):
                                        continue
                                    else:  
                                        # Lista todos os arquivos na pasta
                                        arquivos = os.listdir(aux_path_XML_destino)

                                        for arquivo in arquivos:
                                            caminho_arquivo = os.path.join(aux_path_XML_destino, arquivo)

                                            if os.path.isfile(caminho_arquivo):
                                                try:
                                                    # Abrir o arquivo XML no modo leitura
                                                    with open(caminho_arquivo, 'r'):

                                                        # Comparar cProd do arquivo XML com o valor PN do Excel
                                                        det_element = comparar_cprod(caminho_arquivo, PN, NF)

                                                        if det_element:
                                                            # Buscar o valor de <orig> e <NCM> dentro do mesmo <det>
                                                            CST = det_element.getElementsByTagName('orig')
                                                            ICMS = det_element.getElementsByTagName('pICMS')
                                                            NCM_prod = det_element.getElementsByTagName('NCM')
                                                            Quantidade_prod = det_element.getElementsByTagName('qCom')

                                                            chave_resultado = f"{PN}_{STATUS_planilha}_{NF}"

                                                            if CST and NCM_prod:
                                                                resultados[chave_resultado] = {
                                                                    'VALOR': element_Vprod_value if element_Vprod else None,
                                                                    'CST': str(CST[0].firstChild.data) if CST else None,
                                                                    'ICMS': float(ICMS[0].firstChild.data) if ICMS else None,
                                                                    'NCM': str(NCM_prod[0].firstChild.data) if NCM_prod else None,
                                                                    'QTD.': float(Quantidade_prod[0].firstChild.data) if Quantidade_prod else None
                                                                }

                                                            else:
                                                                log_instance.log_message(f'Elementos CST ou NCM não encontrados no arquivo {arquivo}.')

                                                except AssertionError:
                                                    pass

                                                except Exception as e:
                                                    log_instance.log_message(f'Erro ao processar o arquivo {arquivo}: {e}')

                                            element_peças_pendentes_devolução1 = WebDriverWait(driver, 200).until(
                                            EC.presence_of_element_located((By.XPATH, '(//input[@id="submit"])[4]'))
                                            )

                                    element_peças_pendentes_devolução1.click()

                                except AssertionError:
                                    pass
            
                                except Exception as e:
                                    print()

                                log_instance.log_message(f'{idx_planilha}: Chamado {CHAMADO}/{PN} inserido com sucesso.')
                                idx_planilha += 1

                        except AssertionError:
                            pass

                        except Exception as e:

                            log_instance.log_message(f'{idx_planilha}: Chamado {CHAMADO}/{PN} não consta no site.')
                            idx_planilha += 1

                elif STATUS_planilha.upper() == 'GOOD' or STATUS_planilha.lower() == 'nan':
                    try:
                        if os.path.exists(aux_path_XML_destino):
                            arquivos = os.listdir(aux_path_XML_destino)

                            for arquivo in arquivos:
                                caminho_arquivo = os.path.join(aux_path_XML_destino, arquivo)

                                if os.path.isfile(caminho_arquivo):
                                    det_element = comparar_cprod(caminho_arquivo, PN, NF)

                                    if det_element:
                                        vProd = det_element.getElementsByTagName('vProd')
                                        CST = det_element.getElementsByTagName('orig')
                                        ICMS = det_element.getElementsByTagName('pICMS')
                                        NCM_prod = det_element.getElementsByTagName('NCM')
                                        Quantidade_prod = det_element.getElementsByTagName('qCom')

                                        if CST and NCM_prod and vProd:

                                            chave_resultado = f"{PN}_{STATUS_planilha}_{NF}"
                                            
                                            # Adiciona as informações ao dicionário
                                            resultados[chave_resultado] = {
                                                'VALOR': str(vProd[0].firstChild.data) if vProd else None,
                                                'CST': str(CST[0].firstChild.data) if CST else None,
                                                'ICMS': float(ICMS[0].firstChild.data) if ICMS else None,
                                                'NCM': str(NCM_prod[0].firstChild.data) if NCM_prod else None,
                                                'QTD.': float(Quantidade_prod[0].firstChild.data) if Quantidade_prod else None
                                            }

                                        else:
                                            log_instance.log_message(f'Elementos CST ou NCM não encontrados no arquivo {arquivo}.')
                                            
                        log_instance.log_message(f'{idx_planilha}: Chamado {CHAMADO}/{PN} inserido com sucesso.')
                        idx_planilha += 1

                    except AssertionError:
                        pass
                    
                    except Exception as e:
                        log_instance.log_message(f'Não foi possível abrir o arquivo {idx_planilha}: {e}')

                        idx_planilha += 1

        except:

            driver.quit()

            log_instance.log_message('Erro: Verifique se alguma planilha foi selecionada')

            botao_voltar.disabled = False
            botao_voltar.content.color = ft.Colors.WHITE
            page.update()

        # Criando uma cópia das colunas selecionadas do DataFrame original
        planilha_copia = planilha_df.iloc[:, [0, 1, 2]].copy()

        # Renomeando as colunas da cópia
        planilha_copia.columns = ['CHAMADO', 'REF NF', 'PART NUMBER']

        # Adicionando as colunas ao DataFrame a partir do dicionário
        def tratar(valor):
            return str(valor).replace('.0', '').strip()

        def buscar_resultado(row):
            chave = f"{tratar(row['PN CAIXA.'])}_{tratar(row['STATUS'])}_{tratar(row['NF OU XML PARA BAIXAR (PROGRAMA)'])}"
            return resultados.get(chave, {})

        dados = [buscar_resultado(row) for _, row in planilha_df.iterrows()]

        planilha_copia['VALOR'] = [d.get('VALOR') for d in dados]
        planilha_copia['CST']   = [d.get('CST') for d in dados]
        planilha_copia['ICMS']  = [d.get('ICMS') for d in dados]
        planilha_copia['NCM']   = [d.get('NCM') for d in dados]
        planilha_copia['QTD.']  = [d.get('QTD.') for d in dados]

        # Substituindo pontos por vírgulas na coluna VALOR
        planilha_copia['VALOR'] = planilha_copia['VALOR'].astype(str).str.replace('.', ',', regex = False)

        # Exportando o DataFrame para Excel
        planilha_copia.to_excel('Planilha Dell valores_devolução formatada.xlsx', index = False)

        if os.path.exists(f'{download_dir}\\Planilha Dell valores_devolução formatada.xlsx'):
            os.remove(f'{download_dir}\\Planilha Dell valores_devolução formatada.xlsx')

        shutil.move('Planilha Dell valores_devolução formatada.xlsx', download_dir)

        log_instance.log_message(f'Planilha criada com sucesso: Planilha Dell valores_devolução formatada.xlsx')

        log_instance.log_message('Inserção de valores concluída.')
        
        try:
            time.sleep(3)
            driver.quit()
        except:
            pass

        # Reabilitar o botão "Voltar" após a conclusão
        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE
        page.update()

    def parar(e):
        dialog.open = False
        page.update()
        return
    
    container_aviso = ft.Container(
        content = ft.Column(
            controls = [
                ft.Text(f'Caso existam arquivos na pasta abaixo, serão excluídos: \n\n{aux_path_XML_destino}\n\n Deseja continuar?', text_align = ft.TextAlign.CENTER, color = 'white', size = 15)
            ],
            scroll = ft.ScrollMode.AUTO,
        ),
        margin = ft.margin.only(top = 10),
        height = 180,
        width = 450,
        alignment = ft.alignment.center
    )

    dialog = ft.AlertDialog(
        title = ft.Container(
            content = ft.Row(
                controls = [
                    ft.Icon(ft.Icons.WARNING, color = 'red', size = 30),
                    ft.Text('Aviso', color = 'white', size = 30),
                    ft.Icon(ft.Icons.WARNING, color = 'red', size = 30)],
                alignment = ft.MainAxisAlignment.CENTER
            )
        ),
        content = ft.Column(
            [
                container_aviso,
                ft.Row(
                    [
                        ft.ElevatedButton('Sim', on_click = continuar, bgcolor = 'red', color = 'white', width = 100),
                        ft.ElevatedButton('Não', on_click = parar, bgcolor = 'red', color = 'white', width = 100)
                    ],
                    alignment = ft.MainAxisAlignment.CENTER,
                )
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing = 20,
            width = 450,
            height = 250,
            scroll = ft.ScrollMode.AUTO,
        ),
        bgcolor = 'black',
        actions = [],
    )

    # Adiciona o AlertDialog à página
    page.dialog = dialog

    page.add(dialog)
    dialog.open = True

    # Chama a função de exibir o popup
    mostrar_confirmacao(page)

# Botoes_devolução_HP
def valores_devolução_HP(page: ft.Page, log_instance, views, current_view):

    def mostrar_confirmacao(page):
        dialog.open = True
        page.update()
        
    # Função para o botão "Sim"
    def continuar(e):

        planilha_manager = PlanilhaManager()

        navigate(current_view = current_view, page = page, views = views, view_name = 'logs e informações devolução')

        botao_voltar = views['logs e informações devolução'].content.controls[1].content.controls[0].controls[0]
        botao_voltar.disabled = True
        botao_voltar.content.color = ft.Colors.GREY_400
        dialog.open = False
        page.update()

        # Lista todos os arquivos na pasta
        arquivos_excluir = os.listdir(aux_path_XML_destino)

        log_instance.log_message('Iniciando inserção de valores...')

        for arquivo_excluir in arquivos_excluir:
            caminho_arquivo_excluir = os.path.join(aux_path_XML_destino, arquivo_excluir)
            if os.path.isfile(caminho_arquivo_excluir):
                os.remove(caminho_arquivo_excluir)

        try:
            conectar_email_e_baixar_arquivos_HP(log_instance)
        except:
            log_instance.log_message('Erro: Verifique se alguma planilha foi selecionada. Feche o navegador')
            botao_voltar.disabled = False
            botao_voltar.content.color = ft.Colors.WHITE
            page.update()
            return
        
        dialog.open = False
        page.update()

        idx_planilha = 1

        resultados_hp = {}

        if planilha_manager.is_planilha_atualizada():
            planilha_manager.recarregar_planilha()
        
        planilha_df = planilha_manager.get_planilha()

        try:
            for index, row in planilha_df.iloc[:].iterrows():
                CHAMADO = str(row[0]).replace('.0', '').strip()
                NF = str(row[1]).replace('.0', '').strip()
                PN = str(row[2]).strip()
                STATUS_planilha = str(row[10]).strip()

                try:
                    if os.path.exists(aux_path_XML_destino):

                        arquivos = os.listdir(aux_path_XML_destino)

                        for arquivo in arquivos:
                            caminho_arquivo = os.path.join(aux_path_XML_destino, arquivo)

                            if os.path.isfile(caminho_arquivo):
                                det_element = comparar_cprod(caminho_arquivo, PN, NF)

                                if det_element:
                                        
                                    vProd = det_element.getElementsByTagName('vProd')
                                    CST = det_element.getElementsByTagName('orig')
                                    ICMS = det_element.getElementsByTagName('pICMS')
                                    NCM_prod = det_element.getElementsByTagName('NCM')
                                    Quantidade_prod = det_element.getElementsByTagName('qCom')
                                    
                                    if CST and NCM_prod and vProd:

                                        if STATUS_planilha.upper() == 'GOOD' or STATUS_planilha.lower() == 'nan':

                                            chave_resultado = f"{PN}_{STATUS_planilha}_{NF}"
                                            
                                            # Adiciona as informações ao dicionário
                                            resultados_hp[chave_resultado] = {
                                                'VALOR': str(vProd[0].firstChild.data) if vProd else None,
                                                'CST': str(CST[0].firstChild.data) if CST else None,
                                                'ICMS': float(ICMS[0].firstChild.data) if ICMS else None,
                                                'NCM': str(NCM_prod[0].firstChild.data) if NCM_prod else None,
                                                'QTD.': float(Quantidade_prod[0].firstChild.data) if Quantidade_prod else None
                                            }
                      
                                        elif STATUS_planilha in ['DEFECTIVE', 'DOA']:

                                            chave_resultado = f"{PN}_{STATUS_planilha}_{NF}"
                                    
                                            resultados_hp[chave_resultado] = {
                                                'VALOR': float(f'{float(vProd[0].firstChild.data) / 10:.2f}') if vProd else None,
                                                'CST': str(CST[0].firstChild.data) if CST else None,
                                                'ICMS': float(ICMS[0].firstChild.data) if ICMS else None,
                                                'NCM': str(NCM_prod[0].firstChild.data) if NCM_prod else None,
                                                'QTD.': float(Quantidade_prod[0].firstChild.data) if Quantidade_prod else None
                                            }
                                    else:
                                        log_instance.log_message(f'Elementos CST ou NCM não encontrados no arquivo {arquivo}.')
                                        
                    log_instance.log_message(f'{idx_planilha}: Chamado {CHAMADO}/{PN} inserido com sucesso.')
                    idx_planilha += 1

                except AssertionError:
                    pass
                    
                except Exception as e:
                    log_instance.log_message(f'Não foi possível abrir o arquivo {idx_planilha}: {e}')

                    idx_planilha += 1
        
        except:

            log_instance.log_message('Erro: Verifique se alguma planilha foi selecionada')

            botao_voltar.disabled = False
            botao_voltar.content.color = ft.Colors.WHITE
            page.update()

        # Criando uma cópia das colunas selecionadas do DataFrame original
        planilha_copia = planilha_df.iloc[:, [0, 1, 2]].copy()

        # Renomeando as colunas da cópia
        planilha_copia.columns = ['CHAMADO', 'REF NF', 'PART NUMBER']

        # Adicionando as colunas ao DataFrame a partir do dicionário
        def tratar(valor):
            return str(valor).replace('.0', '').strip()

        def buscar_resultado(row):
            chave = f"{tratar(row['PN CAIXA.'])}_{tratar(row['STATUS'])}_{tratar(row['NF OU XML PARA BAIXAR (PROGRAMA)'])}"
            return resultados_hp.get(chave, {})

        dados = [buscar_resultado(row) for _, row in planilha_df.iterrows()]

        planilha_copia['VALOR'] = [d.get('VALOR') for d in dados]
        planilha_copia['CST']   = [d.get('CST') for d in dados]
        planilha_copia['NCM']   = [d.get('NCM') for d in dados]
        planilha_copia['QTD.']  = [d.get('QTD.') for d in dados]
        planilha_copia['ICMS']  = [d.get('ICMS') for d in dados]

        # Substituindo pontos por vírgulas na coluna VALOR
        planilha_copia['VALOR'] = planilha_copia['VALOR'].astype(str).str.replace('.', ',', regex = False)


        # Exportando o DataFrame para Excel
        planilha_copia.to_excel('Planilha HP valores_devolução formatada.xlsx', index = False)

        if os.path.exists(f'{download_dir}\\Planilha HP valores_devolução formatada.xlsx'):
            os.remove(f'{download_dir}\\Planilha HP valores_devolução formatada.xlsx')

        shutil.move('Planilha HP valores_devolução formatada.xlsx', download_dir)

        log_instance.log_message(f'Planilha criada com sucesso: Planilha HP valores_devolução formatada.xlsx')

        log_instance.log_message('Inserção de valores concluída.')

        # Reabilitar o botão "Voltar" após a conclusão
        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE

        page.update()
            
    def parar(e):
        dialog.open = False
        page.update()
        return
    
    container_aviso = ft.Container(
        content = ft.Column(
            controls = [
                ft.Text(f'Caso existam arquivos na pasta abaixo, serão excluídos: \n\n{aux_path_XML_destino}\n\n Deseja continuar?', text_align = ft.TextAlign.CENTER, color = 'white', size = 15)
            ],
            scroll = ft.ScrollMode.AUTO,
        ),
        margin = ft.margin.only(top = 10),
        height = 180,
        width = 450,
        alignment = ft.alignment.center
    )

    dialog = ft.AlertDialog(
        title = ft.Container(
            content = ft.Row(
                controls = [
                    ft.Icon(ft.Icons.WARNING, color = 'red', size = 30),
                    ft.Text('Aviso', color = 'white', size = 30),
                    ft.Icon(ft.Icons.WARNING, color = 'red', size = 30)],
                alignment = ft.MainAxisAlignment.CENTER
            )
        ),
        content = ft.Column(
            [
                container_aviso,
                ft.Row(
                    [
                        ft.ElevatedButton('Sim', on_click = continuar, bgcolor = 'red', color = 'white', width = 100),
                        ft.ElevatedButton('Não', on_click = parar, bgcolor = 'red', color = 'white', width = 100)
                    ],
                    alignment = ft.MainAxisAlignment.CENTER,
                )
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing = 20,
            width = 450,
            height = 250,
            scroll = ft.ScrollMode.AUTO,
        ),
        bgcolor = 'black',
        actions = [],
    )

    # Adiciona o AlertDialog à página
    page.dialog = dialog

    page.add(dialog)
    dialog.open = True

    # Chama a função de exibir o popup
    mostrar_confirmacao(page)

def criar_etiquetas_devolução_dell(log_instance, page: ft.Page, views, current_view):

    planilha_manager = PlanilhaManager()

    navigate(
        current_view = current_view,
        page = page,
        views = views,
        view_name = 'logs e informações devolução'
    )

    botao_voltar = views['logs e informações devolução'].content.controls[1].content.controls[0].controls[0]
    botao_voltar.disabled = True
    botao_voltar.content.color = ft.Colors.GREY_400
    page.update()

    try:

        if planilha_manager.is_planilha_atualizada():
            planilha_manager.recarregar_planilha()

        planilha_df = planilha_manager.get_planilha()

        documento = Document()

        # Configuração da etiqueta (10,00 x 5,08 cm)
        section = documento.sections[0]
        section.page_width = Cm(10)
        section.page_height = Cm(5.08)
        section.left_margin = Cm(1.19)
        section.right_margin = Cm(1.19)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)

        etiquetas_processadas = 0

        log_instance.log_message(
            f"Iniciando criação das etiquetas de devolução. Total de linhas na planilha: {len(planilha_df)}"
        )

        for index, row in planilha_df.iterrows():

            try:
                CHAMADO = str(row[0]).replace('.0', '').strip()
                NF = str(row[1]).replace('.0', '').strip()
                PN_CAIXA = str(row[2]).strip()
                PN_DEV = str(row[3]).strip()
                NF_SAIDA = str(row[11]).replace('.0', '').strip()
                STATUS_planilha = str(row[10]).strip()

                # Ignorar linhas vazias
                if not CHAMADO or CHAMADO.lower() == 'nan':
                    continue

                # Nova página para cada etiqueta
                if etiquetas_processadas > 0:
                    documento.add_page_break()

                # Cria um parágrafo vazio
                linha = documento.add_paragraph()

                # Função para adicionar borda inferior
                def add_bottom_border(paragraph):
                    p = paragraph._p  # acesso ao XML do parágrafo
                    pPr = p.get_or_add_pPr()
                    pbdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')  # linha simples
                    bottom.set(qn('w:sz'), '6')        # espessura da linha
                    bottom.set(qn('w:space'), '0')     # espaço acima da linha
                    bottom.set(qn('w:color'), '000000') # cor da linha
                    pbdr.append(bottom)
                    pPr.append(pbdr)

                style = documento.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)
                style.paragraph_format.line_spacing = 1.0

                # Primeiro parágrafo: CHAMADO e NF
                paragrafo_topo = documento.add_paragraph()
                paragrafo_topo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_topo = paragrafo_topo.add_run(f'Chamado: {CHAMADO} | NF Entrada: {NF}')
                run_topo.font.name = 'Calibri (Corpo)'
                run_topo.bold = True
                run_topo.font.size = Pt(11)

                # Parágrafo com linha horizontal
                linha = documento.add_paragraph()
                add_bottom_border(linha)  # aqui adiciona a linha

                # Próximo parágrafo: PN CAIXA, PN DEV, STATUS
                paragrafo_dados = documento.add_paragraph()
                paragrafo_dados.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragrafo_dados.add_run(
                    f'\nPart Number Recebido: {PN_CAIXA}\n'
                    f'\nPart Number Devol.: {PN_DEV}\n'
                    f'\nNF Saída: {NF_SAIDA} | Status: {STATUS_planilha}'
                )

                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Calibri (Corpo)'

                etiquetas_processadas += 1

                log_instance.log_message(
                    f"Chamado {CHAMADO} | NF Entrada {NF} - processado com sucesso"
                )

            except Exception as e:
                log_instance.log_message(
                    f"Erro ao processar Chamado {CHAMADO} | NF {NF}: {e}"
                )

        # Salvar somente ao final
        caminho_arquivo = os.path.join(download_dir, "Etiquetas Devolução.docx")

        if os.path.exists(caminho_arquivo):
            os.remove(caminho_arquivo)

        documento.save(caminho_arquivo)

        log_instance.log_message(
            f"Documento criado com sucesso: Etiquetas Devolução.docx ({etiquetas_processadas} etiquetas)"
        )

    except Exception as e:
        log_instance.log_message(f"Erro geral na criação das etiquetas: {e}")

    finally:
        botao_voltar.disabled = False
        botao_voltar.content.color = ft.Colors.WHITE
        page.update()
    